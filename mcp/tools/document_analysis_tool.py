"""
MCP Tool for Document Analysis using python-docx library.

문서 분석 도구:
- python-docx 라이브러리를 사용하여 Word 문서 텍스트 추출
- LLM API 미사용 (라이브러리 기반 분석)
- 지원 포맷: DOCX
- CrewAI BaseTool 패턴 준수
"""
from typing import Type, List
from pathlib import Path
from pydantic import BaseModel, Field, ConfigDict
from crewai.tools import BaseTool

from ai.core.logger import get_logger


class DocumentAnalysisInput(BaseModel):
    """DocumentAnalysisTool 입력 스키마"""

    filepath: str = Field(
        ...,
        description="분석할 Word 문서 파일의 절대 경로 (.docx)"
    )


class DocumentAnalysisCrewAITool(BaseTool):
    """python-docx를 사용한 Word 문서 분석 CrewAI 도구"""

    name: str = "document_analysis"
    description: str = """
    Word 문서 파일(.docx)에서 텍스트를 추출하여 내용을 분석합니다.

    사용 예시:
    - filepath: "/path/to/document.docx"

    지원 포맷: DOCX

    주의: 이 도구는 텍스트 추출만 수행합니다. 추출된 텍스트의 해석이 필요하면
    PlannerAgent에게 위임하세요.
    """
    args_schema: Type[BaseModel] = DocumentAnalysisInput

    model_config = ConfigDict(arbitrary_types_allowed=True)

    def model_post_init(self, __context):
        """Pydantic v2 post-initialization hook for logger setup"""
        super().model_post_init(__context)
        object.__setattr__(self, 'logger', get_logger(__name__))

    def _run(self, filepath: str) -> str:
        """
        Word 문서 파일에서 텍스트 추출 및 분석.

        Args:
            filepath: 문서 파일 경로

        Returns:
            추출된 텍스트 및 문서 메타데이터
        """
        try:
            self.logger.info(f"Starting document analysis: {filepath}")

            # 파일 존재 확인
            file_path = Path(filepath)
            if not file_path.exists():
                error_msg = f"문서 파일을 찾을 수 없습니다: {filepath}"
                self.logger.error(error_msg)
                return f"❌ {error_msg}"

            # 파일 확장자 확인
            if file_path.suffix.lower() != ".docx":
                error_msg = f"지원하지 않는 문서 포맷: {file_path.suffix} (DOCX만 지원)"
                self.logger.error(error_msg)
                return f"❌ {error_msg}"

            # 파일 크기 제한 확인 (Fix #6)
            MAX_FILE_SIZE_MB = 50
            file_size_mb = file_path.stat().st_size / (1024 * 1024)
            if file_size_mb > MAX_FILE_SIZE_MB:
                error_msg = f"문서 파일이 너무 큽니다: {file_size_mb:.1f}MB (최대: {MAX_FILE_SIZE_MB}MB)"
                self.logger.error(error_msg)
                return f"❌ {error_msg}"

            # python-docx로 문서 분석
            result = self._extract_text_from_docx(file_path)

            self.logger.info(f"Document analysis completed: {len(result)} characters")
            return result

        except Exception as exc:
            error_msg = f"문서 분석 실패: {exc}"
            self.logger.exception(error_msg)
            return f"❌ {error_msg}"

    def _extract_text_from_docx(self, file_path: Path) -> str:
        """
        python-docx를 사용하여 DOCX 파일에서 텍스트 추출.

        Args:
            file_path: 문서 파일 경로

        Returns:
            추출된 텍스트 및 메타데이터 (포맷팅된 문자열)
        """
        try:
            # python-docx 임포트 (느슨한 의존성)
            try:
                from docx import Document
            except ImportError as exc:
                raise ImportError(
                    "python-docx 패키지가 설치되지 않았습니다. "
                    "'pip install python-docx' 후 다시 시도하세요."
                ) from exc

            # 문서 로드
            self.logger.debug(f"Loading document: {file_path}")
            doc = Document(file_path)

            # 메타데이터 추출
            core_properties = doc.core_properties
            metadata = {
                "제목": core_properties.title or "(제목 없음)",
                "작성자": core_properties.author or "(작성자 미상)",
                "생성일": str(core_properties.created) if core_properties.created else "(날짜 미상)",
                "단락 수": len(doc.paragraphs),
                "파일 크기": f"{file_path.stat().st_size / 1024:.1f} KB"
            }

            # 텍스트 추출 (모든 단락)
            paragraphs: List[str] = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:  # 빈 단락 제외
                    paragraphs.append(text)

            # 결과 포맷팅
            result_lines = ["📄 Word 문서 분석 결과\n"]

            # 메타데이터 섹션
            result_lines.append("### 📋 문서 정보")
            for key, value in metadata.items():
                result_lines.append(f"- {key}: {value}")

            # 텍스트 내용 섹션
            result_lines.append("\n### 📝 문서 내용")
            if paragraphs:
                # 단락 구분하여 출력
                for idx, para in enumerate(paragraphs, 1):
                    # 너무 긴 단락은 요약 (1000자 제한)
                    if len(para) > 1000:
                        para_text = para[:1000] + "... (이하 생략)"
                    else:
                        para_text = para
                    result_lines.append(f"\n[단락 {idx}]")
                    result_lines.append(para_text)
            else:
                result_lines.append("(문서에 텍스트 내용이 없습니다)")

            # 통계 정보
            total_chars = sum(len(p) for p in paragraphs)
            result_lines.append(f"\n### 📊 통계")
            result_lines.append(f"- 총 단락 수: {len(paragraphs)}개")
            result_lines.append(f"- 총 문자 수: {total_chars:,}자")

            result_text = "\n".join(result_lines)
            self.logger.debug(f"Extracted {len(paragraphs)} paragraphs, {total_chars} characters")

            return result_text

        except Exception as exc:
            error_msg = f"DOCX 텍스트 추출 실패: {exc}"
            self.logger.exception(error_msg)
            raise RuntimeError(error_msg) from exc
